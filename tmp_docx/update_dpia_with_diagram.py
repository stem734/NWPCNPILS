from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import math

SRC = Path('/Users/steve/Downloads/MyMedInfo DPIA Policy v1.3 completed - DPO DCB advice.docx')
OUT = Path('/Users/steve/Downloads/MyMedInfo DPIA Policy v1.3 final with data flow diagram.docx')
PNG = Path('/Users/steve/Downloads/MyMedInfo DPIA data flow diagram.png')

W, H = 1800, 1100
img = Image.new('RGB', (W, H), 'white')
d = ImageDraw.Draw(img)
try:
    font_title = ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial Bold.ttf', 42)
    font_box = ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial.ttf', 28)
    font_box_bold = ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial Bold.ttf', 30)
    font_small = ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial.ttf', 22)
except Exception:
    font_title = ImageFont.load_default()
    font_box = ImageFont.load_default()
    font_box_bold = ImageFont.load_default()
    font_small = ImageFont.load_default()

BLUE = '#005EB8'
LIGHT = '#EAF4FF'
GREEN = '#007F3B'
GREEN_TINT = '#EEF8F0'
GREY = '#425563'
BORDER = '#AEB7BD'
TEXT = '#1F2933'


def wrap(text, font, width):
    words = text.split()
    lines = []
    cur = ''
    for w in words:
        test = (cur + ' ' + w).strip()
        if d.textbbox((0, 0), test, font=font)[2] <= width:
            cur = test
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def box(x, y, w, h, title, body, fill, title_fill=None):
    d.rounded_rectangle((x, y, x + w, y + h), radius=22, fill=fill, outline=BORDER, width=3)
    if title_fill:
        d.rounded_rectangle((x + 18, y + 16, x + w - 18, y + 64), radius=12, fill=title_fill)
        d.text((x + 34, y + 24), title, font=font_box_bold, fill='white')
        body_y = y + 88
    else:
        d.text((x + 24, y + 22), title, font=font_box_bold, fill=BLUE)
        body_y = y + 70
    for i, line in enumerate(wrap(body, font_box, w - 48)):
        d.text((x + 24, body_y + i * 34), line, font=font_box, fill=TEXT)


def arrow(x1, y1, x2, y2, label=None, color=GREY):
    d.line((x1, y1, x2, y2), fill=color, width=6)
    ang = math.atan2(y2 - y1, x2 - x1)
    size = 18
    p1 = (x2, y2)
    p2 = (x2 - size * math.cos(ang - math.pi / 6), y2 - size * math.sin(ang - math.pi / 6))
    p3 = (x2 - size * math.cos(ang + math.pi / 6), y2 - size * math.sin(ang + math.pi / 6))
    d.polygon([p1, p2, p3], fill=color)
    if label:
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        bbox = d.textbbox((0, 0), label, font=font_small)
        tw = bbox[2] - bbox[0]
        d.rounded_rectangle((mx - tw / 2 - 12, my - 20, mx + tw / 2 + 12, my + 20), radius=10, fill='white', outline=BORDER)
        d.text((mx - tw / 2, my - 12), label, font=font_small, fill=GREY)

# Title and boundary labels
d.text((70, 40), 'MyMedInfo DPIA Data Flow Diagram', font=font_title, fill=BLUE)
d.text((70, 90), 'Single shared DPIA covering participating GP practices and related NHS controllers using the same service model', font=font_small, fill=GREY)

# Boxes
box(80, 180, 360, 150, 'Participating GP Practice', 'Controller. Selects content, generates patient links or QR codes, and decides use within local care pathways.', LIGHT, BLUE)
box(80, 430, 360, 150, 'Practice / Admin Users', 'Authorised users manage medication cards, templates, practices, and user access through the MyMedInfo dashboards.', LIGHT, BLUE)
box(540, 180, 360, 150, 'MyMedInfo Patient View', 'Patient opens the link, MyMedInfo resolves the permitted card content, and trusted guidance is displayed.', LIGHT, BLUE)
box(540, 430, 360, 150, 'MyMedInfo Dashboards', 'Authenticated web app for content management, practice configuration, and service administration.', LIGHT, BLUE)
box(1020, 120, 320, 180, 'Supabase Auth + Edge Functions', 'Authentication, save/read logic, template resolution, medication card lookup, and audit-related server-side operations.', GREEN_TINT, GREEN)
box(1020, 380, 320, 180, 'Supabase Database', 'Stores medication cards, templates, practice settings, limited audit data, usage counts, and ratings.', GREEN_TINT, GREEN)
box(1420, 180, 280, 150, 'Patient Browser / Device', 'Receives the patient link or QR code and opens the MyMedInfo patient page over HTTPS.', LIGHT, BLUE)
box(1420, 430, 280, 150, 'Approved External Resources', 'Optional patient click-through to NHS or approved third-party guidance linked from the displayed card.', LIGHT, BLUE)

# Arrows
arrow(440, 255, 540, 255, 'Patient link / QR')
arrow(260, 330, 260, 430, 'Content admin')
arrow(440, 505, 540, 505, 'Manage content')
arrow(900, 255, 1020, 220, 'Resolve content')
arrow(900, 505, 1020, 470, 'Authenticated save')
arrow(1180, 300, 1180, 380, 'Read / write')
arrow(1340, 255, 1420, 255, 'HTTPS')
arrow(900, 255, 1420, 255, 'Displayed in browser')
arrow(900, 505, 1420, 505, 'Optional linked resources')
arrow(1580, 330, 1580, 430, 'User clicks')

# Footer notes
d.rounded_rectangle((80, 830, 1720, 1030), radius=18, fill='#F7FAFC', outline=BORDER, width=2)
d.text((110, 860), 'Controller interpretation for this DPIA', font=font_box_bold, fill=BLUE)
notes = [
    'This DPIA covers the common MyMedInfo processing activity as a single shared assessment across participating GP practices acting as separate data controllers.',
    'Each controller remains responsible for lawful basis, local transparency, local governance, and the decision to use MyMedInfo in its own organisation.',
    'The platform and common safeguards are assessed once here to avoid duplication and keep risk and mitigation review consistent across all participating controllers.',
]
for i, note in enumerate(notes):
    d.text((120, 910 + i * 34), u'• ' + note, font=font_small, fill=TEXT)

img.save(PNG)

# DOCX edit helpers

def insert_paragraph_after(paragraph, text='', style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def set_cell_text(cell, text):
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.size = Pt(10)


doc = Document(SRC)

# Formalise the controller wording in the title / intro area.
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip() == '[Participating GP Practice] / Nottingham West PCN':
                set_cell_text(cell, 'Participating GP Practices / Nottingham West PCN')
            if 'GP data / work undertaken by [Participating GP Practice]' in cell.text:
                set_cell_text(cell, 'Shared GP practice processing activity. This single DPIA covers all participating GP practices and related NHS organisations acting as separate data controllers using the MyMedInfo service model.')
            if cell.text.strip().startswith('Data Controller: the GP practice sending the patient link.'):
                set_cell_text(cell, 'Data Controllers: each participating GP practice using MyMedInfo. Service operator / processor: Nottingham West PCN / MyMedInfo. Infrastructure sub-processors: Supabase and Vercel.')

# Add formal shared-DPIA wording into the introduction.
intro_anchor = None
for p in doc.paragraphs:
    if 'This DPIA is to be completed where the practice is the Data Controller' in p.text:
        intro_anchor = p
        break
if intro_anchor is not None:
    p1 = insert_paragraph_after(intro_anchor, 'This DPIA has been completed as a single shared assessment for the MyMedInfo service across participating GP practices and related NHS organisations acting as separate data controllers. It assesses the common data flows, system functions, categories of personal data, risks, and technical and organisational controls that apply across the shared service model.', style='No Spacing')
    insert_paragraph_after(p1, 'Each participating data controller remains responsible for its own lawful basis, local privacy information, local governance decisions, and the accuracy and appropriateness of the content it provides or approves. This shared DPIA is intended to avoid duplication while ensuring a consistent assessment of risk and mitigation across all participating controllers.', style='No Spacing')

# Add a formal data flow diagram and explanatory text after the data flow introduction.
flow_anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith('The next two boxes describe the data flows involved.'):
        flow_anchor = p
        break
if flow_anchor is not None:
    heading = insert_paragraph_after(flow_anchor, 'Shared Service Data Flow Diagram', style='Heading 2')
    explainer = insert_paragraph_after(heading, 'The diagram below summarises the common MyMedInfo data flows covered by this single DPIA. It should be read together with the inbound and outbound flow tables and adopted by each participating controller as the shared processing view for the service.', style='No Spacing')
    image_para = insert_paragraph_after(explainer, '', style='Normal')
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = image_para.add_run()
    run.add_picture(str(PNG), width=Inches(6.8))
    caption = insert_paragraph_after(image_para, 'Figure 1. MyMedInfo shared DPIA data flow diagram', style='No Spacing')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    follow = insert_paragraph_after(caption, 'This figure reflects the standard MyMedInfo operating model. Where a participating controller introduces material local variation, that organisation should review whether a local addendum is needed alongside this shared DPIA.', style='No Spacing')

# Strengthen data flow narrative placeholders if present.
for p in doc.paragraphs:
    if p.text.strip() == 'a.':
        prev = p._p.getprevious()
        if prev is not None and 'The next two boxes describe the data flows involved' in ''.join(t.text for t in []):
            pass

# Save
if OUT.exists():
    OUT.unlink()
doc.save(OUT)
print(str(OUT))
print(str(PNG))
