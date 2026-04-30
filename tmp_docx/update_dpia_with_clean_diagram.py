from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont
import math

SRC = Path('/Users/steve/Downloads/MyMedInfo DPIA Policy v1.3 completed - DPO DCB advice.docx')
OUT = Path('/Users/steve/Downloads/MyMedInfo DPIA Policy v1.3 final with data flow diagram.docx')
PNG = Path('/Users/steve/Downloads/MyMedInfo DPIA data flow diagram.png')

W, H = 1900, 1200
img = Image.new('RGB', (W, H), 'white')
d = ImageDraw.Draw(img)
try:
    font_title = ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial Bold.ttf', 44)
    font_sub = ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial.ttf', 24)
    font_head = ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial Bold.ttf', 24)
    font_body = ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial.ttf', 24)
    font_note = ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial.ttf', 21)
except Exception:
    font_title = font_sub = font_head = font_body = font_note = ImageFont.load_default()

BLUE = '#005EB8'; BLUE_TINT = '#EAF4FF'; GREEN = '#007F3B'; GREEN_TINT = '#EEF8F0'; GREY = '#425563'; TEXT = '#1F2933'; BORDER = '#AEB7BD'; SOFT = '#F7FAFC'

def wrap(text, font, width):
    words=text.split(); lines=[]; cur=''
    for word in words:
        trial=(cur+' '+word).strip()
        if d.textbbox((0,0), trial, font=font)[2] <= width:
            cur=trial
        else:
            if cur: lines.append(cur)
            cur=word
    if cur: lines.append(cur)
    return lines

def draw_box(x,y,w,h,title,body,fill,accent):
    d.rounded_rectangle((x,y,x+w,y+h), radius=22, fill=fill, outline=BORDER, width=3)
    d.rounded_rectangle((x+18,y+18,x+w-18,y+62), radius=12, fill=accent)
    for i,line in enumerate(wrap(title,font_head,w-70)):
        d.text((x+30,y+23+i*24), line, font=font_head, fill='white')
    yy = y+88
    for line in wrap(body,font_body,w-48):
        d.text((x+24,yy), line, font=font_body, fill=TEXT)
        yy += 32

def label(x,y,text):
    bbox=d.textbbox((0,0),text,font=font_note); tw=bbox[2]-bbox[0]; th=bbox[3]-bbox[1]
    d.rounded_rectangle((x,y,x+tw+22,y+th+14), radius=10, fill='white', outline=BORDER, width=2)
    d.text((x+11,y+7), text, font=font_note, fill=GREY)

def arrow(x1,y1,x2,y2,color=GREY):
    d.line((x1,y1,x2,y2), fill=color, width=6)
    ang=math.atan2(y2-y1,x2-x1); size=16
    p1=(x2,y2); p2=(x2-size*math.cos(ang-math.pi/6), y2-size*math.sin(ang-math.pi/6)); p3=(x2-size*math.cos(ang+math.pi/6), y2-size*math.sin(ang+math.pi/6))
    d.polygon([p1,p2,p3], fill=color)

d.text((70,38),'MyMedInfo DPIA Data Flow Diagram',font=font_title,fill=BLUE)
d.text((70,92),'Single shared DPIA covering participating GP practices and related NHS controllers using the same service model',font=font_sub,fill=GREY)

boxes={
 'gp':(90,180,360,185), 'admin':(90,430,360,200), 'patient':(560,180,380,185), 'dash':(560,430,380,200),
 'auth':(1060,160,390,185), 'db':(1060,420,390,210), 'browser':(1520,180,290,185), 'ext':(1520,430,290,200)
}

draw_box(*boxes['gp'],'Participating GP Practice','Controller. Selects content, generates patient links or QR codes, and decides local use within the care pathway.',BLUE_TINT,BLUE)
draw_box(*boxes['admin'],'Practice / Admin Users','Authorised users manage medication cards, templates, practices, and user access through the MyMedInfo dashboards.',BLUE_TINT,BLUE)
draw_box(*boxes['patient'],'MyMedInfo Patient View','Patient opens the link. MyMedInfo resolves the permitted card content and displays trusted guidance.',BLUE_TINT,BLUE)
draw_box(*boxes['dash'],'MyMedInfo Dashboards','Authenticated web app for content management, practice configuration, and service administration.',BLUE_TINT,BLUE)
draw_box(*boxes['auth'],'Supabase Auth + Edge Functions','Authentication plus server-side save, read, template resolution, medication lookup, and audit logic.',GREEN_TINT,GREEN)
draw_box(*boxes['db'],'Supabase Database','Stores medication cards, templates, practice settings, limited audit data, usage counts, and ratings.',GREEN_TINT,GREEN)
draw_box(*boxes['browser'],'Patient Browser / Device','Receives the patient link or QR code and opens the MyMedInfo patient page over HTTPS.',BLUE_TINT,BLUE)
draw_box(*boxes['ext'],'Approved External Resources','Optional patient click-through to NHS or approved third-party guidance linked from the displayed card.',BLUE_TINT,BLUE)

arrow(450,272,560,272); label(432,226,'Patient link / QR')
arrow(940,272,1060,252); label(920,220,'Resolve content')
arrow(1450,272,1520,272); label(1460,220,'HTTPS')
arrow(450,530,560,530); label(432,484,'Manage content')
arrow(940,530,1060,520); label(930,484,'Authenticated save')
arrow(1255,345,1255,420); label(1190,356,'Read / write')
arrow(940,315,1520,315); label(1320,286,'Displayed in browser')
arrow(1665,365,1665,430); label(1578,378,'User clicks')
arrow(1450,530,1520,530); label(1210,506,'Optional linked resources')

d.rounded_rectangle((90,840,1810,1090), radius=22, fill=SOFT, outline=BORDER, width=3)
d.text((120,875),'Controller interpretation for this DPIA',font=font_head,fill=BLUE)
notes=[
 'This DPIA covers the common MyMedInfo processing activity as a single shared assessment across participating GP practices acting as separate data controllers.',
 'Each controller remains responsible for lawful basis, local transparency, local governance, and the decision to use MyMedInfo in its own organisation.',
 'The platform and common safeguards are assessed once here to avoid duplication and keep risk and mitigation review consistent across all participating controllers.'
]
y=930
for note in notes:
    for i,line in enumerate(wrap('• '+note,font_note,1620)):
        d.text((130,y+i*28),line,font=font_note,fill=TEXT)
    y += 62

img.save(PNG)

def insert_after(paragraph,text='',style=None):
    new_p=OxmlElement('w:p'); paragraph._p.addnext(new_p)
    para=paragraph._parent.add_paragraph(); para._p.getparent().remove(para._p); para._p=new_p
    if style: para.style=style
    if text: para.add_run(text)
    return para

def set_cell(cell,text):
    cell.text=text
    for p in cell.paragraphs:
        for r in p.runs: r.font.size=Pt(10)

doc=Document(SRC)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            txt=cell.text.strip()
            if txt == '[Participating GP Practice] / Nottingham West PCN':
                set_cell(cell,'Participating GP Practices / Nottingham West PCN')
            elif 'GP data / work undertaken by [Participating GP Practice]' in txt:
                set_cell(cell,'Shared GP practice processing activity. This single DPIA covers all participating GP practices and related NHS organisations acting as separate data controllers using the MyMedInfo service model.')
            elif txt.startswith('Data Controller: the GP practice sending the patient link.'):
                set_cell(cell,'Data Controllers: each participating GP practice using MyMedInfo. Service operator / processor: Nottingham West PCN / MyMedInfo. Infrastructure sub-processors: Supabase and Vercel.')
intro_anchor=next((p for p in doc.paragraphs if 'This DPIA is to be completed where the practice is the Data Controller' in p.text),None)
if intro_anchor:
    p1=insert_after(intro_anchor,'This DPIA has been completed as a single shared assessment for the MyMedInfo service across participating GP practices and related NHS organisations acting as separate data controllers. It assesses the common data flows, system functions, categories of personal data, risks, and technical and organisational controls that apply across the shared service model.','No Spacing')
    insert_after(p1,'Each participating data controller remains responsible for its own lawful basis, local privacy information, local governance decisions, and the accuracy and appropriateness of the content it provides or approves. This shared DPIA is intended to avoid duplication while ensuring a consistent assessment of risk and mitigation across all participating controllers.','No Spacing')
flow_anchor=next((p for p in doc.paragraphs if p.text.strip().startswith('The next two boxes describe the data flows involved.')),None)
if flow_anchor:
    p2=insert_after(flow_anchor,'Shared Service Data Flow Diagram','Heading 2')
    p3=insert_after(p2,'The diagram below summarises the common MyMedInfo data flows covered by this single DPIA. It should be read together with the inbound and outbound flow tables and adopted by each participating controller as the shared processing view for the service.','No Spacing')
    p4=insert_after(p3,'','Normal'); p4.alignment=WD_ALIGN_PARAGRAPH.CENTER; p4.add_run().add_picture(str(PNG),width=Inches(6.9))
    p5=insert_after(p4,'Figure 1. MyMedInfo shared DPIA data flow diagram','No Spacing'); p5.alignment=WD_ALIGN_PARAGRAPH.CENTER
    insert_after(p5,'This figure reflects the standard MyMedInfo operating model. Where a participating controller introduces material local variation, that organisation should review whether a local addendum is needed alongside this shared DPIA.','No Spacing')
if OUT.exists(): OUT.unlink()
doc.save(OUT)
print(OUT)
print(PNG)
